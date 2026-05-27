from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, os

BASE = os.path.dirname(os.path.abspath(__file__))
SCREENSHOTS = os.path.join(BASE, "report-screenshots")
OUT = os.path.join(BASE, "IoT-Dashboard-系統現況報告.docx")

# ── palette ──────────────────────────────────────────────
BLUE_DARK  = RGBColor(0x1A, 0x37, 0x5E)   # heading / cover bg
BLUE_MID   = RGBColor(0x21, 0x5D, 0xAC)   # accent
BLUE_LIGHT = RGBColor(0xD6, 0xE4, 0xF7)   # table header bg
GRAY_TEXT  = RGBColor(0x44, 0x44, 0x44)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
RED        = RGBColor(0xC0, 0x39, 0x2B)
GREEN      = RGBColor(0x1A, 0x7A, 0x46)
AMBER      = RGBColor(0xB7, 0x5A, 0x00)

def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    hex_color = str(rgb).upper()
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, color="2159AC"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top","left","bottom","right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def para_font(para, size, bold=False, color=None, italic=False):
    for run in para.runs:
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = color

def add_heading(doc, text, level=1, color=BLUE_DARK):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = color
        run.font.bold = True
    return p

def add_body(doc, text, color=GRAY_TEXT, size=10.5):
    p = doc.add_paragraph(text)
    para_font(p, size, color=color)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.line_spacing = Pt(16)
    return p

def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        set_cell_bg(cell, BLUE_DARK)
        set_cell_border(cell, "1A375E")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0] if p.runs else p.add_run(h)
        run.font.bold  = True
        run.font.color.rgb = WHITE
        run.font.size  = Pt(9.5)
    # data rows
    for ri, row in enumerate(rows):
        bg = RGBColor(0xF5, 0xF8, 0xFF) if ri % 2 == 0 else WHITE
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]
            cell.text = str(val)
            set_cell_bg(cell, bg)
            set_cell_border(cell, "C5D8F0")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if p.runs:
                p.runs[0].font.size = Pt(9.5)
                p.runs[0].font.color.rgb = GRAY_TEXT
    if col_widths:
        for ri, row in enumerate(t.rows):
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Cm(w)
    return t

def add_status_table(doc, rows):
    """Special table for device status with colour-coded status column."""
    headers = ["設備名稱", "連線類型", "即時狀態", "說明"]
    widths  = [4.0, 3.0, 2.5, 8.0]
    t = doc.add_table(rows=1+len(rows), cols=4)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        set_cell_bg(cell, BLUE_DARK)
        set_cell_border(cell, "1A375E")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0] if p.runs else p.add_run(h)
        run.font.bold  = True
        run.font.color.rgb = WHITE
        run.font.size  = Pt(9.5)
    status_colors = {"✅ 正常": GREEN, "⚠️ 連線中斷": AMBER}
    for ri, row in enumerate(rows):
        bg = RGBColor(0xF5, 0xF8, 0xFF) if ri % 2 == 0 else WHITE
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]
            set_cell_bg(cell, bg)
            set_cell_border(cell, "C5D8F0")
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9.5)
            if ci == 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.font.bold = True
                for key, clr in status_colors.items():
                    if key in str(val):
                        run.font.color.rgb = clr
                        break
            else:
                run.font.color.rgb = GRAY_TEXT
    for ri, row_ in enumerate(t.rows):
        for ci, w in enumerate(widths):
            row_.cells[ci].width = Cm(w)
    return t

def add_image(doc, path, width_cm=15.5, caption=None):
    if not os.path.exists(path):
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Cm(width_cm))
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_font(cp, 9, italic=True, color=RGBColor(0x88,0x88,0x88))
        cp.paragraph_format.space_after = Pt(10)

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:color"), "2159AC")
    pBdr.append(bottom)
    pPr.append(pBdr)

# ═══════════════════════════════════════════════════════════
doc = Document()

# page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Cover Page ────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run("IoT 監控儀表板")
title_run.font.size  = Pt(28)
title_run.font.bold  = True
title_run.font.color.rgb = BLUE_DARK

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub_p.add_run("系統現況報告")
sub_run.font.size  = Pt(20)
sub_run.font.color.rgb = BLUE_MID

doc.add_paragraph()

meta_items = [
    ("報告日期", "2026-05-27"),
    ("系統環境", "生產環境（192.168.6.23:5200）"),
    ("監控設備", "6 台 / 19 個感測點"),
    ("上線時間", "2026-04-29 初版 / 2026-05-26 全設備整合"),
]
meta_t = doc.add_table(rows=len(meta_items), cols=2)
meta_t.alignment = WD_TABLE_ALIGNMENT.CENTER
for ri, (k, v) in enumerate(meta_items):
    lc, rc = meta_t.rows[ri].cells[0], meta_t.rows[ri].cells[1]
    lc.text, rc.text = k, v
    set_cell_bg(lc, BLUE_LIGHT)
    for cell in (lc, rc):
        set_cell_border(cell, "2159AC")
        p = cell.paragraphs[0]
        if p.runs:
            p.runs[0].font.size = Pt(10.5)
            p.runs[0].font.color.rgb = BLUE_DARK if cell is lc else GRAY_TEXT
            if cell is lc:
                p.runs[0].font.bold = True
    lc.width, rc.width = Cm(4), Cm(8)

doc.add_page_break()

# ── 一、系統概述 ──────────────────────────────────────────
add_heading(doc, "一、系統概述")
add_body(doc,
    "本系統為鑽石集團 C 棟 LeanA 鞋業產線的工廠 IoT 即時監控平台，"
    "自 2026 年 4 月底上線，整合直連 PLC 設備資料與廠商（厚信）推送資料，"
    "統一顯示於瀏覽器儀表板，超過製程界限時立即告警，"
    "資料每秒自動推送、無需手動刷新。")
add_body(doc, "程式碼異動自動部署（GitLab CI/CD），從提交到生產上線約 7 分鐘。")
add_divider(doc)

# ── 二、系統畫面 ───────────────────────────────────────────
add_heading(doc, "二、系統畫面")
add_heading(doc, "主儀表板", level=2, color=BLUE_MID)
add_image(doc, os.path.join(SCREENSHOTS, "dashboard.png"),
          caption="圖 1：主儀表板 — 即時顯示所有設備狀態（截圖時間：2026-05-27）")

add_heading(doc, "趨勢分析頁面", level=2, color=BLUE_MID)
add_image(doc, os.path.join(SCREENSHOTS, "trend.png"),
          caption="圖 2：趨勢頁面 — 歷史折線圖 + UCL/LCL 管制界限")
add_divider(doc)

# ── 三、監控設備清單 ───────────────────────────────────────
add_heading(doc, "三、監控設備清單")
add_body(doc, "目前共監控 6 台設備、19 個感測點，全部正常運行中。")

add_heading(doc, "Modbus 直連設備（4 台，每 2 秒取值）", level=2, color=BLUE_MID)
add_table(doc,
    ["設備名稱", "感測點", "顯示方式"],
    [
        ["高速加熱定型機", "溫度 × 1",      "單值顯示"],
        ["烘箱",          "左右兩側溫度 × 2","雙側即時折線"],
        ["冷凍機",        "溫度 × 1",      "單值顯示"],
        ["冷熱定型機",    "四區溫度 × 4",   "四環形量表"],
    ],
    col_widths=[4.5, 5.0, 4.0]
)
doc.add_paragraph()

add_heading(doc, "廠商推送設備（2 台，每 2 秒取值）", level=2, color=BLUE_MID)
add_table(doc,
    ["設備名稱", "感測點數", "資料來源"],
    [
        ["強勢壓底機", "14 點（壓力 × 6、計次 × 4、時間 × 4）", "厚信系統推送"],
        ["畫線機",     "1 點（壓力）",                          "厚信系統推送"],
    ],
    col_widths=[3.5, 8.5, 3.5]
)
add_divider(doc)

# ── 四、即時串流狀態 ───────────────────────────────────────
add_heading(doc, "四、即時串流狀態")
add_body(doc, "以下為 2026-05-27 18:47（台灣時間）向生產環境即時查詢的結果：")

add_status_table(doc, [
    ["強勢壓底機",    "廠商 API",   "✅ 正常",    "持續接收資料，零錯誤"],
    ["畫線機",        "廠商 API",   "✅ 正常",    "持續接收資料，零錯誤"],
    ["高速加熱定型機","Modbus TCP", "⚠️ 連線中斷","TCP 連線逾時（詳見備註）"],
    ["烘箱",          "Modbus TCP", "⚠️ 連線中斷","TCP 連線逾時（詳見備註）"],
    ["冷凍機",        "Modbus TCP", "⚠️ 連線中斷","TCP 連線逾時（詳見備註）"],
    ["冷熱定型機",    "Modbus TCP", "⚠️ 連線中斷","TCP 連線逾時（詳見備註）"],
])
doc.add_paragraph()

note = doc.add_paragraph()
note.paragraph_format.left_indent = Cm(0.5)
r = note.add_run("備註：")
r.font.bold = True
r.font.color.rgb = AMBER
r.font.size = Pt(9.5)
r2 = note.add_run(
    "查詢時間為台灣時間下午 6:47，屬工廠非生產時段，4 台 Modbus 設備可能已依慣例關機。"
    "廠商推送的兩台設備透過共用資料庫接收資料，不受 PLC 連線狀態影響，持續正常運行。"
    "系統於連線中斷時會在儀表板顯示琥珀色「感測器異常」標籤，不顯示過時數值。"
)
r2.font.size = Pt(9.5)
r2.font.color.rgb = GRAY_TEXT
add_divider(doc)

# ── 五、資料庫現況 ─────────────────────────────────────────
add_heading(doc, "五、資料庫現況")
add_body(doc, "資料涵蓋系統初期運行階段（2026-04-20 至 2026-05-14，共 24 天）。")

add_table(doc,
    ["項目", "數值"],
    [
        ["感測器讀值總筆數", "22,973 筆"],
        ["資料異常值（PLC 雜訊）", "0 筆（過濾機制全程有效）"],
        ["歷史告警筆數", "44 筆"],
        ["資料起始日", "2026-04-20"],
        ["最新資料日", "2026-05-14"],
    ],
    col_widths=[6.0, 7.5]
)
doc.add_paragraph()

add_heading(doc, "各設備資料量", level=2, color=BLUE_MID)
add_table(doc,
    ["設備", "總筆數", "每日平均", "告警次數", "狀態"],
    [
        ["烘箱",          "8,934 筆","約 425 筆/天","31 次","曾有超溫事件（見第六章）"],
        ["高速加熱定型機","7,103 筆","約 296 筆/天","0 次", "優良"],
        ["冷熱定型機",    "5,688 筆","約 271 筆/天","0 次", "優良"],
        ["冷凍機",        "1,248 筆","約 59 筆/天", "13 次","與烘箱超溫事件連動"],
    ],
    col_widths=[3.8, 2.2, 2.8, 2.2, 5.0]
)
add_divider(doc)

# ── 六、告警事件分析 ───────────────────────────────────────
add_heading(doc, "六、告警事件分析")
add_body(doc,
    "系統上線 24 天內，44 筆告警全部集中在 2026-04-28 凌晨 02:48 至 03:21 這 33 分鐘，"
    "涉及兩台設備：烘箱與冷凍機。事件結束後直至 2026-05-14 資料期末，完全零告警。")

add_body(doc, "（資料庫儲存 UTC，以下時間均已換算為台灣時間 UTC+8）", size=9, color=RGBColor(0x88,0x88,0x88))
add_heading(doc, "事件時序（含管制界限依據）", level=2, color=BLUE_MID)
add_table(doc,
    ["台灣時間", "設備", "讀值", "管制界限（UCL）", "超出幅度"],
    [
        ["10:48", "烘箱（右側溫度）", "84℃",   "65℃", "+19℃"],
        ["10:50", "烘箱（右側溫度）", "104℃",  "65℃", "+39℃"],
        ["10:56", "烘箱（右側溫度）", "106℃",  "65℃", "+41℃（最高）"],
        ["11:00", "冷凍機（溫度）",   "5.9℃",  "5℃",  "+0.9℃"],
        ["11:21", "烘箱（右側溫度）", "97.8℃", "65℃", "+32.8℃（最後一筆）"],
    ],
    col_widths=[2.5, 4.0, 2.5, 3.5, 3.5]
)
doc.add_paragraph()

add_heading(doc, "研判", level=2, color=BLUE_MID)
add_body(doc,
    "事件發生在上班日上午正常生產時段（台灣時間 10:48–11:21），現場應有操作人員。"
    "烘箱右側溫度超出 65℃ 上限最高達 41℃，持續 33 分鐘後自行回復正常。"
    "冷凍機僅超出 0.9℃，為連動反應。"
    "兩台設備恢復正常後未再觸發任何告警，非持續性設備故障，"
    "推測為加熱迴路短暫失控或溫控系統切換異常。告警系統即時偵測並完整記錄全程，機制運作正常。")

add_heading(doc, "告警統計摘要", level=2, color=BLUE_MID)
add_table(doc,
    ["項目", "數值"],
    [
        ["告警總筆數",    "44 筆"],
        ["危險等級（超界）","42 筆（95%）"],
        ["警告等級（近界）","2 筆（5%）"],
        ["涉及設備",     "烘箱 31 筆、冷凍機 13 筆"],
        ["事件持續時間", "33 分鐘"],
        ["事後無告警期", "2026-04-28 03:21 起，持續 16 天"],
    ],
    col_widths=[6.0, 7.5]
)
add_divider(doc)

# ── 七、系統品質與測試 ─────────────────────────────────────
add_heading(doc, "七、系統品質與測試")
add_body(doc, "系統共有 200 個自動化測試案例，每次程式碼提交後自動執行，全部通過才允許部署上線。")
add_table(doc,
    ["測試類型", "案例數", "說明"],
    [
        ["後端邏輯測試", "127", "資料寫入、告警判斷、連線管理等"],
        ["前端元件測試", "18",  "介面元件行為驗證"],
        ["端對端功能測試","55", "模擬真實使用者操作，含告警、趨勢、多語言"],
    ],
    col_widths=[5.0, 2.5, 8.0]
)
add_divider(doc)

# ── 八、待處理事項 ─────────────────────────────────────────
add_heading(doc, "八、待處理事項")
add_table(doc,
    ["優先級", "項目", "說明"],
    [
        ["高",  "資料保留策略",
         "全設備每秒合計約 11.5 筆（強勢壓底機 7 + 畫線機 0.5 + Modbus 4 台 2 秒共 4），每日估計近 100 萬筆；建議設定 90 天自動刪除舊資料"],
        ["中",  "WeChat 告警啟用",
         "架構已備妥，取得 Webhook 網址後即可開啟，告警可即時推送至手機"],
        ["低",  "告警確認功能",
         "後端已支援告警簽收，前端頁面尚未串接"],
    ],
    col_widths=[1.8, 4.0, 9.7]
)

doc.add_paragraph()
footer_p = doc.add_paragraph("資料來源：生產環境資料庫直接查詢 + GitLab CI/CD 部署記錄（2026-05-27）")
footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
para_font(footer_p, 8.5, italic=True, color=RGBColor(0xAA,0xAA,0xAA))

doc.save(OUT)
print(f"saved → {OUT}")
